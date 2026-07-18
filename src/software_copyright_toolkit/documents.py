from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from software_copyright_toolkit import (
    SkippedFile,
    count_supported_file,
    iter_files,
    read_text_lossless,
)


LINES_PER_PAGE = 50
FRONT_BACK_PAGES = 30
FRONT_BACK_LINES = LINES_PER_PAGE * FRONT_BACK_PAGES
IDENTIFICATION_TOTAL_LINES = FRONT_BACK_LINES * 2
MAX_WORD_LINE_CHARS = 92
MODULE_END_PATTERNS = (
    re.compile(r"^\s*[}\])]+[;,\])}]*\s*$"),
    re.compile(r"^\s*};?\s*$"),
    re.compile(r"^\s*end\s*[;,%#]*\s*$", re.IGNORECASE),
    re.compile(r"^\s*end(?:function|class|methods|properties|for|if|while|switch)\b.*$", re.IGNORECASE),
)


@dataclass(frozen=True)
class GeneratedDocuments:
    identification_docx: Path | None
    full_docx: Path | None
    source_line_count: int
    source_content_line_count: int
    identification_line_count: int
    identification_content_line_count: int


def sanitize_filename(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", value).strip()
    cleaned = re.sub(r"\s+", "_", cleaned)
    return cleaned or "未命名软件"


@dataclass(frozen=True)
class SourceLine:
    text: str
    counts_as_source: bool
    is_source_text: bool


def collect_source_lines(target: Path) -> tuple[list[SourceLine], list[SkippedFile]]:
    lines: list[SourceLine] = []
    skipped: list[SkippedFile] = []

    for path in sorted(iter_files(target), key=lambda item: str(item.relative_to(target)).lower()):
        try:
            text = read_text_lossless(path)
        except (OSError, UnicodeError) as exc:
            skipped.append(SkippedFile(path, f"读取失败或非文本文件: {exc}"))
            continue

        if count_supported_file(path, text) is None:
            skipped.append(SkippedFile(path, f"未适配的文件后缀: {path.suffix.lower() or '(无后缀)'}"))
            continue

        rel = str(path.relative_to(target)).replace("\\", "/")
        lines.append(SourceLine(f"===== 文件: {rel} =====", counts_as_source=False, is_source_text=False))
        lines.extend(
            SourceLine(line, counts_as_source=True, is_source_text=True)
            for line in text.splitlines()
            if line.strip()
        )

    return lines, skipped


def generate_identification_documents(
    target: Path,
    output_dir: Path,
    software_name: str,
    version: str,
    include_identification: bool = True,
    include_full: bool = True,
) -> GeneratedDocuments:
    source_lines, _ = collect_source_lines(target)
    if not source_lines:
        raise ValueError("没有可用于生成程序鉴别材料的已适配代码文件。")
    if not include_identification and not include_full:
        raise ValueError("请至少选择一种 Word 材料。")

    name_part = sanitize_filename(software_name)
    version_part = sanitize_filename(version)
    prefix = f"{name_part}_{version_part}_程序鉴别材料"

    identification_path = output_dir / f"{prefix}_前后30页.docx" if include_identification else None
    full_path = output_dir / f"{prefix}_全量代码.docx" if include_full else None
    source_content_line_count = count_source_content_lines(source_lines)
    display_lines = wrap_source_lines(source_lines)
    identification_lines = select_identification_lines(display_lines, source_content_line_count)
    identification_content_line_count = count_source_content_lines(identification_lines)

    if identification_path is not None:
        write_code_docx(
            identification_path,
            identification_lines,
            software_name=software_name,
            version=version,
            page_break_interval=LINES_PER_PAGE
            if len(identification_lines) >= IDENTIFICATION_TOTAL_LINES
            else None,
        )
    if full_path is not None:
        write_code_docx(
            full_path,
            display_lines,
            software_name=software_name,
            version=version,
        )

    return GeneratedDocuments(
        identification_docx=identification_path,
        full_docx=full_path,
        source_line_count=len(display_lines),
        source_content_line_count=source_content_line_count,
        identification_line_count=len(identification_lines),
        identification_content_line_count=identification_content_line_count,
    )


def wrap_source_lines(lines: list[SourceLine]) -> list[SourceLine]:
    wrapped: list[SourceLine] = []
    for line in lines:
        if len(line.text) <= MAX_WORD_LINE_CHARS:
            wrapped.append(line)
            continue

        parts = split_long_line(line.text, MAX_WORD_LINE_CHARS)
        for index, part in enumerate(parts):
            wrapped.append(
                SourceLine(
                    part,
                    counts_as_source=line.counts_as_source and index == 0,
                    is_source_text=line.is_source_text,
                )
            )
    return wrapped


def split_long_line(line: str, width: int) -> list[str]:
    return [line[index : index + width] for index in range(0, len(line), width)] or [line]


def select_identification_lines(lines: list[SourceLine], source_content_line_count: int) -> list[SourceLine]:
    if source_content_line_count < IDENTIFICATION_TOTAL_LINES:
        return lines

    if len(lines) < IDENTIFICATION_TOTAL_LINES:
        return lines

    first_part_end = FRONT_BACK_LINES
    end_index = find_module_end_index(lines, minimum_index=IDENTIFICATION_TOTAL_LINES)
    if end_index is None:
        return lines

    back_start = end_index - FRONT_BACK_LINES
    if back_start < first_part_end:
        return lines

    selected = lines[:first_part_end] + lines[back_start:end_index]
    if not is_valid_identification_selection(selected):
        return lines
    return selected


def count_source_content_lines(lines: list[SourceLine]) -> int:
    return sum(1 for line in lines if line.counts_as_source)


def find_module_end_index(lines: list[SourceLine], minimum_index: int) -> int | None:
    for index in range(len(lines), minimum_index - 1, -1):
        line = lines[index - 1]
        if line.is_source_text and is_module_end_line(line.text):
            return index
    return None


def is_module_end_line(line: str) -> bool:
    return any(pattern.match(line) for pattern in MODULE_END_PATTERNS)


def is_valid_identification_selection(lines: list[SourceLine]) -> bool:
    return (
        len(lines) == IDENTIFICATION_TOTAL_LINES
        and lines[-1].is_source_text
        and is_module_end_line(lines[-1].text)
    )


def write_code_docx(
    path: Path,
    lines: list[SourceLine],
    software_name: str,
    version: str,
    page_break_interval: int | None = None,
) -> None:
    doc = Document()
    setup_document(doc, software_name, version)

    for index, line in enumerate(lines, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(14)
        run = paragraph.add_run(line.text if line.text else " ")
        apply_run_font(run)
        if page_break_interval and index % page_break_interval == 0 and index < len(lines):
            run.add_break(WD_BREAK.PAGE)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def setup_document(doc: Document, software_name: str, version: str) -> None:
    enable_field_updates(doc)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    setup_header(section, software_name, version)
    setup_footer(section)


def enable_field_updates(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def setup_header(section, software_name: str, version: str) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(17), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.paragraph_format.space_after = Pt(0)

    left_run = paragraph.add_run(f"{software_name} {version}")
    apply_run_font(left_run)
    tab_run = paragraph.add_run("\t")
    apply_run_font(tab_run)
    add_field(paragraph, "PAGE")


def setup_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    add_field(paragraph, "PAGE")
    slash_run = paragraph.add_run(" / ")
    apply_run_font(slash_run)
    add_field(paragraph, "NUMPAGES")


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def apply_run_font(run) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    apply_run_font(run)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)
